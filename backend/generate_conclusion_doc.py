from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN TƯƠNG LAI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Kết Luận
doc.add_heading('1. Kết luận', level=1)
p1 = doc.add_paragraph('Qua quá trình nghiên cứu, thiết kế và phát triển, dự án "Chi Tiêu Thông Minh" đã hoàn thành tốt các mục tiêu đề ra ban đầu, xây dựng thành công một nền tảng hỗ trợ tài chính cá nhân toàn diện. Cụ thể, dự án đã đạt được những kết quả đáng chú ý sau:')

ul1 = doc.add_paragraph(style='List Bullet')
ul1.add_run('Về mặt công nghệ: ').bold = True
ul1.add_run('Áp dụng thành công các công nghệ hiện đại nhất hiện nay với kiến trúc tách biệt Frontend (Next.js 14, React 19, Tailwind CSS) và Backend (.NET 8, MongoDB). Sự kết hợp này mang lại giao diện mượt mà (Gamified UI), trải nghiệm người dùng cao cấp cùng khả năng xử lý dữ liệu lớn ở tốc độ cao.')

ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run('Về mặt nghiệp vụ: ').bold = True
ul2.add_run('Hệ thống đã thu thập và đồng bộ thành công kho dữ liệu khổng lồ với gần 400 thẻ tín dụng từ hàng chục ngân hàng lớn tại Việt Nam. Thuật toán gợi ý thẻ hoạt động chính xác dựa trên mức thu nhập, chi tiêu và danh mục ưu tiên của người dùng.')

ul3 = doc.add_paragraph(style='List Bullet')
ul3.add_run('Tích hợp Trí tuệ nhân tạo (AI): ').bold = True
ul3.add_run('Chatbot tư vấn tài chính được tích hợp thuật toán phân loại Intent (ID3) giúp phân tích nhu cầu người dùng qua các câu hội thoại tự nhiên, qua đó nâng tầm trải nghiệm tương tác số.')

# 2. Kiến nghị
doc.add_heading('2. Hướng phát triển và Kiến nghị', level=1)
p2 = doc.add_paragraph('Mặc dù hệ thống đã đáp ứng tốt các nghiệp vụ cốt lõi, nhưng để nền tảng "Chi Tiêu Thông Minh" thực sự trở thành một siêu ứng dụng tài chính (Super App) có sức cạnh tranh trên thị trường, nhóm phát triển đề xuất một số hướng mở rộng trong tương lai:')

# Khuyến nghị 1
h1 = doc.add_heading('2.1. Nâng cấp AI Chatbot tích hợp Mô hình Ngôn ngữ Lớn (LLM)', level=2)
doc.add_paragraph('Thay thế công cụ phân loại Intent hiện tại bằng các mô hình AI tiên tiến (như OpenAI GPT-4 hoặc Google Gemini) để Chatbot có khả năng ghi nhớ ngữ cảnh hội thoại, tự động tóm tắt chi tiêu và đưa ra các lời khuyên tài chính cá nhân hóa sâu sắc hơn.')

# Khuyến nghị 2
h2 = doc.add_heading('2.2. Mở rộng Hệ sinh thái trên nền tảng Di động (Mobile App)', level=2)
doc.add_paragraph('Phát triển thêm phiên bản ứng dụng di động (sử dụng React Native hoặc Flutter) tích hợp công nghệ định vị (GPS). Khi người dùng đi ngang qua các nhà hàng, trung tâm thương mại có liên kết hoàn tiền với thẻ tín dụng họ đang sở hữu, ứng dụng sẽ gửi thông báo đẩy (Push Notification) theo thời gian thực.')

# Khuyến nghị 3
h3 = doc.add_heading('2.3. Tích hợp định danh điện tử (e-KYC) và Liên kết Ngân hàng', level=2)
doc.add_paragraph('Bắt tay với các đối tác ngân hàng và tổ chức tài chính để tích hợp công nghệ định danh điện tử (e-KYC). Điều này cho phép người dùng đăng ký mở thẻ tín dụng ngay trên nền tảng "Chi Tiêu Thông Minh" thay vì phải chuyển hướng sang website của ngân hàng, giúp rút ngắn phễu chuyển đổi (Conversion Funnel).')

# Khuyến nghị 4
h4 = doc.add_heading('2.4. Phát triển Dashboard Quản lý Dòng tiền (Personal Finance Management)', level=2)
doc.add_paragraph('Cho phép người dùng liên kết các tài khoản thẻ tín dụng thực tế vào ứng dụng (thông qua Open Banking API). Từ đó cung cấp các biểu đồ thống kê thói quen chi tiêu, cảnh báo rủi ro khi sắp vượt hạn mức hoặc nhắc nhở tự động ngày đến hạn thanh toán dư nợ (Due Date) nhằm tránh bị phạt lãi suất.')

doc.save('/Users/sivan/chitieuthongminh/KetLuan_KienNghi_ChiTieuThongMinh.docx')
print("Đã tạo file Kết Luận và Kiến Nghị DOCX thành công!")
