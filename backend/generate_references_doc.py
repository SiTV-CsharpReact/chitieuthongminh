from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DANH MỤC TÀI LIỆU THAM KHẢO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Vietnamese References
doc.add_heading('1. Tài liệu chuyên ngành và Nghiệp vụ', level=1)

p1 = doc.add_paragraph(style='List Number')
p1.add_run('Ngân hàng Nhà nước Việt Nam. ').bold = True
p1.add_run('Các quy định, thông tư về hoạt động phát hành và sử dụng thẻ tín dụng ngân hàng tại Việt Nam.')

p2 = doc.add_paragraph(style='List Number')
p2.add_run('Website chính thức của các Ngân hàng Thương mại: ').bold = True
p2.add_run('Tài liệu biểu phí, điều khoản sử dụng, chính sách hoàn tiền (Cashback), điểm thưởng (Rewards) được công bố công khai từ các ngân hàng: VIB, VPBank, Techcombank, HSBC, Shinhan Bank...')

p3 = doc.add_paragraph(style='List Number')
p3.add_run('TS. Nguyễn Văn A (nếu có sách tham khảo thì thay tên, hoặc để chung chung): ').bold = True
p3.add_run('Giáo trình Quản trị Tài chính Cá nhân và Các công cụ Tín dụng thời đại số.')

# Foreign References
doc.add_heading('2. Tài liệu Công nghệ và Nền tảng (Tiếng Anh)', level=1)

p4 = doc.add_paragraph(style='List Number')
p4.add_run('Microsoft Corporation. ').bold = True
p4.add_run('ASP.NET Core 8.0 Documentation & Entity Framework / MongoDB C# Driver. Truy cập tại: https://learn.microsoft.com/')

p5 = doc.add_paragraph(style='List Number')
p5.add_run('Vercel Inc. ').bold = True
p5.add_run('Next.js 14 Documentation (App Router Architecture). Truy cập tại: https://nextjs.org/docs')

p6 = doc.add_paragraph(style='List Number')
p6.add_run('Meta Platforms, Inc. ').bold = True
p6.add_run('React 18/19 Official Documentation - UI State Management. Truy cập tại: https://react.dev/')

p7 = doc.add_paragraph(style='List Number')
p7.add_run('MongoDB, Inc. ').bold = True
p7.add_run('MongoDB Manual & NoSQL Database Architecture. Truy cập tại: https://www.mongodb.com/docs/')

p8 = doc.add_paragraph(style='List Number')
p8.add_run('Quinlan, J. R. (1986). ').bold = True
p8.add_run('Induction of Decision Trees (Cơ sở lý thuyết thuật toán ID3 áp dụng cho AI Chatbot Intent Classification). Machine Learning, 1(1), 81-106.')

p9 = doc.add_paragraph(style='List Number')
p9.add_run('Tailwind Labs. ').bold = True
p9.add_run('Tailwind CSS Framework Documentation. Truy cập tại: https://tailwindcss.com/docs')

doc.save('/Users/sivan/chitieuthongminh/TaiLieuThamKhao_ChiTieuThongMinh.docx')
print("Đã tạo file Danh Mục Tài Liệu Tham Khảo DOCX thành công!")
