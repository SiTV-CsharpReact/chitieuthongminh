from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('KẾ HOẠCH BẢO MẬT VÀ KỊCH BẢN KIỂM THỬ HỆ THỐNG', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dưới đây là nội dung chi tiết về Các Biện pháp Bảo mật và Kịch bản Kiểm thử (Test Cases) được thiết kế bám sát vào kiến trúc thực tế của hệ thống (Next.js + .NET 8 + MongoDB + AI Chatbot).')

# Section 1
doc.add_heading('PHẦN 1: KẾ HOẠCH VÀ BIỆN PHÁP BẢO MẬT HỆ THỐNG', level=1)
doc.add_paragraph('Dự án áp dụng các tiêu chuẩn bảo mật đa lớp (Multi-layer Security) từ Frontend, Backend đến Database để đảm bảo an toàn dữ liệu và tính toàn vẹn của hệ thống.')

doc.add_heading('1.1. Bảo mật Xác thực & Phân quyền (Authentication & Authorization)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('JSON Web Token (JWT): ').bold = True
p.add_run('Quá trình đăng nhập trả về JWT Token được ký mã hóa bằng khóa bí mật. Token có thời hạn (Expiration) để tự động vô hiệu hóa các phiên đăng nhập cũ, giảm rủi ro bị đánh cắp phiên.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Mã hóa mật khẩu: ').bold = True
p.add_run('Mật khẩu người dùng được băm (Hash) một chiều trước khi lưu vào cơ sở dữ liệu để đảm bảo dù lộ Database cũng không lộ mật khẩu gốc.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Role-Based Access Control (RBAC): ').bold = True
p.add_run('Hệ thống phân tách rõ quyền hạn. Các API nhạy cảm (như thêm, sửa, xóa thẻ tín dụng, cấu hình danh mục) được bảo vệ phân quyền Admin, trong khi API gợi ý thẻ được truy cập bởi User/Guest.')

doc.add_heading('1.2. Bảo mật API & Dữ liệu (Backend Security)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Input Validation & Data Sanitization: ').bold = True
p.add_run('Tất cả dữ liệu đầu vào từ người dùng (đặc biệt ở form Chatbot và Bộ lọc tìm kiếm) đều được kiểm tra tính hợp lệ qua Data Annotations để ngăn chặn XSS (Cross-Site Scripting) và NoSQL Injection.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('CORS (Cross-Origin Resource Sharing): ').bold = True
p.add_run('Backend .NET được cấu hình chỉ chấp nhận các request từ domain/port hợp lệ của Frontend, chặn hoàn toàn các truy cập trái phép từ bên thứ ba.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Bảo vệ AI Engine: ').bold = True
p.add_run('API gọi AI Chatbot được đóng gói ngầm trong Backend thay vì gọi trực tiếp từ Frontend để bảo mật thuật toán nhận diện và quy trình xử lý dữ liệu.')

doc.add_heading('1.3. Bảo mật Cơ sở dữ liệu (Database Security)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Bảo mật Connection String: ').bold = True
p.add_run('Chuỗi kết nối tới cơ sở dữ liệu được tách riêng khỏi mã nguồn, quản lý qua môi trường triển khai (Environment Variables), không lưu trữ trực tiếp trên Git.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Cơ chế sao lưu (Backup): ').bold = True
p.add_run('Hệ thống áp dụng các cấu hình sao lưu định kỳ của cơ sở dữ liệu để đảm bảo an toàn cho dữ liệu thẻ, nhật ký chat và thông tin người dùng.')

# Section 2
doc.add_heading('PHẦN 2: KỊCH BẢN KIỂM THỬ HỆ THỐNG (TEST CASES)', level=1)
doc.add_paragraph('Hoạt động kiểm thử được chia làm 4 nhóm chính để đảm bảo hệ thống vận hành ổn định trước khi đưa vào thực tế.')

doc.add_heading('2.1. Kiểm thử Chức năng cốt lõi (Functional Testing)', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Chức năng'
hdr_cells[2].text = 'Kịch bản kiểm thử (Test Scenario)'
hdr_cells[3].text = 'Kết quả mong đợi'

test_cases_1 = [
    ('TC_01', 'Đăng ký & Đăng nhập', 'Nhập sai tài khoản/mật khẩu hoặc email chưa đúng định dạng.', 'Hệ thống chặn lại, hiển thị thông báo lỗi thân thiện. Không xảy ra gián đoạn.'),
    ('TC_02', 'Thuật toán Gợi ý thẻ', 'Nhập mức chi tiêu 10 triệu cho danh mục "Ăn uống".', 'Hiển thị Gamified Best Card (Thẻ lựa chọn số 1) là thẻ có tỉ lệ hoàn tiền cao nhất cho Ăn uống. Thông tin số tiền hoàn phải được tính đúng.'),
    ('TC_03', 'AI Chatbot Intent', 'Nhắn tin: "Tôi muốn tìm thẻ du lịch".', 'Chatbot nhận diện đúng Intent là "Tìm thẻ du lịch" và trả về danh sách thẻ có lợi ích du lịch.'),
    ('TC_04', 'AI Chatbot Fallback', 'Nhắn tin rác: "asdjaskjd" hoặc hỏi sai chủ đề.', 'Chatbot trả về câu trả lời mặc định, định hướng người dùng hỏi lại về thẻ tín dụng.'),
    ('TC_05', 'Phân trang (Pagination)', 'Tải danh mục có nhiều thẻ với thiết lập kích thước trang là 12.', 'Hệ thống hiển thị đúng số trang, nút Next/Prev hoạt động, thay đổi bộ lọc tự động quay về Trang 1.')
]

for id, func, scenario, expected in test_cases_1:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = func
    row_cells[2].text = scenario
    row_cells[3].text = expected

doc.add_heading('2.2. Kiểm thử Tích hợp (Integration Testing)', level=2)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'ID'
hdr_cells2[1].text = 'Chức năng'
hdr_cells2[2].text = 'Kịch bản kiểm thử (Test Scenario)'
hdr_cells2[3].text = 'Kết quả mong đợi'

test_cases_2 = [
    ('TC_06', 'Đồng bộ Dữ liệu', 'Kiểm tra luồng cập nhật dữ liệu thẻ từ Backend vào Database.', 'Dữ liệu được cập nhật chính xác (Tên thẻ, Ảnh, Quy tắc hoàn tiền) vào cơ sở dữ liệu mà không bị trùng lặp.'),
    ('TC_07', 'JWT Expiration', 'Để màn hình Frontend không thao tác cho đến khi JWT hết hạn, sau đó bấm gọi API.', 'Backend trả về lỗi 401 Unauthorized, Frontend bắt lỗi và đẩy người dùng ra màn hình Đăng nhập.')
]

for id, func, scenario, expected in test_cases_2:
    row_cells = table2.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = func
    row_cells[2].text = scenario
    row_cells[3].text = expected

doc.add_heading('2.3. Kiểm thử Giao diện (UI/UX & Responsive Testing)', level=2)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'ID'
hdr_cells3[1].text = 'Chức năng'
hdr_cells3[2].text = 'Kịch bản kiểm thử (Test Scenario)'
hdr_cells3[3].text = 'Kết quả mong đợi'

test_cases_3 = [
    ('TC_08', 'Dark/Light Mode', 'Chuyển đổi công tắc Theme (Light -> Dark).', 'Toàn bộ giao diện (kể cả khung Chatbot và thẻ Gamified) thay đổi màu nền/chữ đồng bộ, không bị lỗi tương phản.'),
    ('TC_09', 'Responsive Mobile', 'Thu nhỏ màn hình xuống kích thước điện thoại di động.', 'Bố cục lưới thẻ chuyển từ 3-4 cột về 1 cột. Sidebar ẩn vào Hamburger menu. Bảng so sánh hiển thị tốt dạng dọc.'),
    ('TC_10', 'Gamified UI', 'Truy cập trang Gợi ý có dữ liệu.', 'Thẻ top 1 được làm nổi bật với hiệu ứng Glowing/Pulse, Badge "Lựa chọn số 1", tách biệt với danh sách thẻ thường.')
]

for id, func, scenario, expected in test_cases_3:
    row_cells = table3.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = func
    row_cells[2].text = scenario
    row_cells[3].text = expected

doc.add_heading('2.4. Kiểm thử Hiệu năng (Performance Testing)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('TC_11 (Xử lý danh sách lớn): ').bold = True
p.add_run('Đảm bảo API danh sách thẻ phản hồi dưới 500ms khi Database chứa hàng trăm bản ghi nhờ vào index của MongoDB. Giao diện Frontend không bị chậm do đã áp dụng phân trang (Pagination).')

p = doc.add_paragraph(style='List Bullet')
p.add_run('TC_12 (Quản lý Bộ nhớ): ').bold = True
p.add_run('Giám sát tiến trình backend khi xử lý lượng dữ liệu lớn để đảm bảo RAM không bị tràn, tối ưu cấu hình thu gom rác (Garbage Collection) tránh crash server.')

doc.save('/Users/sivan/chitieuthongminh/BaoMat_KiemThu_ChiTieuThongMinh.docx')
print("Đã tạo file DOCX thành công!")
