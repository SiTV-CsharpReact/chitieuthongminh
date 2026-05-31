from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Helper: set all borders on a table
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

# Helper: make header row bold
def set_header_row(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

# Helper: add a database table section
def add_db_table(doc, table_number, title, caption, columns, rows):
    # Section heading
    heading = doc.add_heading(f'{title}', level=3)
    heading.style.font.name = 'Times New Roman'
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.bold = True
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(13)
    cap_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Create table
    num_cols = len(columns)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Column widths: TT narrow, Mô tả wide
    col_widths = [Cm(1.2), Cm(3.5), Cm(3.3), Cm(8.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Header row
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(columns):
        header_cells[i].text = col_name
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, value in enumerate(row_data):
            row_cells[col_idx].text = str(value)
            for paragraph in row_cells[col_idx].paragraphs:
                # Center align TT column
                if col_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()  # spacing


# ===== DOCUMENT CONTENT =====

# Title
title = doc.add_heading('Tài liệu Cơ sở dữ liệu (Database Documentation)', level=1)
for run in title.runs:
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

intro = doc.add_paragraph(
    'Hệ thống sử dụng MongoDB làm cơ sở dữ liệu NoSQL. '
    'Dưới đây là mô tả chi tiết cấu trúc các Collection (Bảng) trong hệ thống.'
)
for run in intro.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

# ---- TABLE DEFINITIONS ----

columns = ['TT', 'Tên cột', 'Kiểu dữ liệu', 'Mô tả']

# 1. User
add_db_table(doc, 1,
    '2.3.1. Bảng người dùng (users)',
    'Bảng 2.1. Bảng users',
    columns,
    [
        [1, 'Id', 'ObjectId', 'Khóa chính tự sinh của MongoDB'],
        [2, 'Email', 'string', 'Địa chỉ email dùng để đăng nhập'],
        [3, 'PasswordHash', 'string', 'Mật khẩu đã được mã hóa (hash)'],
        [4, 'Name', 'string', 'Tên hiển thị của người dùng'],
        [5, 'Avatar', 'string', 'Đường dẫn ảnh đại diện'],
        [6, 'Role', 'string', 'Vai trò người dùng (Admin/User)'],
        [7, 'Provider', 'string', 'Nhà cung cấp đăng nhập (Local, Google, Facebook)'],
        [8, 'ProviderId', 'string', 'ID định danh từ nhà cung cấp bên thứ 3'],
    ]
)

# 2. CreditCard
add_db_table(doc, 2,
    '2.3.2. Bảng thẻ tín dụng (credit_cards)',
    'Bảng 2.2. Bảng credit_cards',
    columns,
    [
        [1,  'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2,  'Name', 'string', 'Tên thẻ tín dụng'],
        [3,  'Bank', 'string', 'Mã ngân hàng phát hành'],
        [4,  'BankName', 'string', 'Tên đầy đủ của ngân hàng'],
        [5,  'BankLogo', 'string', 'Đường dẫn logo ngân hàng'],
        [6,  'ImageUrl', 'string', 'Hình ảnh thẻ tín dụng'],
        [7,  'Link', 'string', 'Đường dẫn trang chi tiết thẻ'],
        [8,  'RegisterUrl', 'string', 'Đường dẫn đăng ký mở thẻ'],
        [9,  'AnnualFee', 'decimal', 'Phí thường niên (VNĐ)'],
        [10, 'CashbackRules', 'array<object>', 'Danh sách quy tắc hoàn tiền (nested)'],
        [11, 'MinSalary', 'decimal', 'Lương tối thiểu yêu cầu (0 = không yêu cầu)'],
        [12, 'Description', 'string', 'Mô tả chi tiết thẻ'],
        [13, 'Benefits', 'array<string>', 'Danh sách đặc quyền/ưu đãi'],
        [14, 'CreditLimit', 'string', 'Hạn mức tín dụng'],
        [15, 'InterestRate', 'string', 'Lãi suất'],
        [16, 'TermsPdfUrl', 'string', 'Link tài liệu điều khoản sử dụng'],
    ]
)

# 2b. CashbackRule (nested)
add_db_table(doc, 2.1,
    '2.3.3. Cấu trúc đối tượng nhúng CashbackRule (trong credit_cards)',
    'Bảng 2.3. Cấu trúc CashbackRule',
    columns,
    [
        [1, 'Category', 'string', 'Danh mục hoàn tiền (Ẩm thực, Mua sắm, Tất cả)'],
        [2, 'Percentage', 'decimal', 'Phần trăm hoàn tiền (%)'],
        [3, 'CapAmount', 'decimal (nullable)', 'Hạn mức hoàn tiền tối đa theo tháng/quý'],
    ]
)

# 3. SpendingData
add_db_table(doc, 3,
    '2.3.4. Bảng dữ liệu chi tiêu (spending_data)',
    'Bảng 2.4. Bảng spending_data',
    columns,
    [
        [1,  'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2,  'UserId', 'string (FK)', 'Khóa ngoại liên kết tới users.Id'],
        [3,  'Amount', 'decimal', 'Số tiền chi tiêu'],
        [4,  'Salary', 'decimal', 'Mức lương hàng tháng của người dùng'],
        [5,  'Category', 'string', 'Danh mục chi tiêu'],
        [6,  'Date', 'DateTime', 'Ngày tháng phát sinh giao dịch'],
        [7,  'Description', 'string', 'Mô tả chi tiết giao dịch'],
        [8,  'IncomeLevel', 'string', 'Mức thu nhập (Thấp, Trung bình, Cao)'],
        [9,  'SpendingHabit', 'string', 'Thói quen chi tiêu (Tiết kiệm, Vừa phải, Phung phí)'],
        [10, 'CreditScoreRange', 'string', 'Dải điểm tín dụng (Poor, Fair, Good, Excellent)'],
        [11, 'RecommendedCardType', 'string', 'Loại thẻ được hệ thống gợi ý'],
    ]
)

# 4. CardPromotion
add_db_table(doc, 4,
    '2.3.5. Bảng khuyến mãi thẻ (card_promotions)',
    'Bảng 2.5. Bảng card_promotions',
    columns,
    [
        [1,  'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2,  'Title', 'string', 'Tiêu đề chương trình khuyến mãi'],
        [3,  'Description', 'string', 'Nội dung mô tả chi tiết'],
        [4,  'ImageUrl', 'string', 'Hình ảnh banner khuyến mãi'],
        [5,  'DiscountRate', 'string', 'Mức giảm giá/hoàn tiền'],
        [6,  'CategoryTab', 'string', 'Nhóm danh mục (Ẩm thực, Mua sắm...)'],
        [7,  'SourceUrl', 'string', 'Link gốc chương trình'],
        [8,  'StartDate', 'string', 'Ngày bắt đầu'],
        [9,  'ValidUntil', 'string', 'Ngày kết thúc / Hạn sử dụng'],
        [10, 'ApplicableCards', 'array<string>', 'Danh sách mã/tên thẻ được áp dụng'],
        [11, 'CreatedAt', 'DateTime', 'Thời gian tạo'],
        [12, 'UpdatedAt', 'DateTime', 'Thời gian cập nhật'],
    ]
)

# 5. Category
add_db_table(doc, 5,
    '2.3.6. Bảng danh mục chi tiêu (categories)',
    'Bảng 2.6. Bảng categories',
    columns,
    [
        [1, 'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2, 'Name', 'string', 'Tên danh mục (Ăn uống, Y tế, ...)'],
        [3, 'Color', 'string', 'Mã màu Hex (mặc định: #3b82f6)'],
        [4, 'Icon', 'string', 'Biểu tượng đại diện cho danh mục'],
        [5, 'MccCodes', 'array<string>', 'Danh sách mã MCC (Merchant Category Code)'],
        [6, 'IsFrequent', 'bool', 'Đánh dấu danh mục phổ biến'],
    ]
)

# 6. Article
add_db_table(doc, 6,
    '2.3.7. Bảng bài viết (articles)',
    'Bảng 2.7. Bảng articles',
    columns,
    [
        [1,  'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2,  'Title', 'string', 'Tiêu đề bài viết'],
        [3,  'Slug', 'string', 'Đường dẫn thân thiện SEO'],
        [4,  'Excerpt', 'string', 'Đoạn trích dẫn ngắn'],
        [5,  'Content', 'string', 'Nội dung chi tiết (HTML/Markdown)'],
        [6,  'Category', 'string', 'Chuyên mục bài viết'],
        [7,  'Author', 'string', 'Tên tác giả (Mặc định: Admin)'],
        [8,  'CoverImage', 'string', 'Ảnh bìa bài viết'],
        [9,  'CreatedAt', 'DateTime', 'Thời gian tạo'],
        [10, 'UpdatedAt', 'DateTime', 'Thời gian cập nhật'],
    ]
)

# 7. ArticleCategory
add_db_table(doc, 7,
    '2.3.8. Bảng thể loại bài viết (article_categories)',
    'Bảng 2.8. Bảng article_categories',
    columns,
    [
        [1, 'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2, 'Name', 'string', 'Tên chuyên mục'],
        [3, 'Slug', 'string', 'Đường dẫn thân thiện SEO'],
        [4, 'Description', 'string', 'Mô tả chuyên mục'],
        [5, 'Color', 'string', 'Mã màu hiển thị giao diện'],
        [6, 'CreatedAt', 'DateTime', 'Thời gian tạo'],
        [7, 'UpdatedAt', 'DateTime', 'Thời gian cập nhật'],
    ]
)

# 8. ChatLog (Chatbot AI)
add_db_table(doc, 8,
    '2.3.9. Bảng lịch sử hội thoại Chatbot AI (chat_logs)',
    'Bảng 2.9. Bảng chat_logs',
    columns,
    [
        [1, 'Id', 'ObjectId', 'Khóa chính tự sinh'],
        [2, 'SessionId', 'string', 'Mã phiên hội thoại (mỗi lần mở chat)'],
        [3, 'UserId', 'string (FK)', 'Khóa ngoại liên kết tới users.Id (nếu đã đăng nhập)'],
        [4, 'Message', 'string', 'Nội dung tin nhắn của người dùng'],
        [5, 'Reply', 'string', 'Nội dung phản hồi của Chatbot AI'],
        [6, 'Intent', 'string', 'Ý định được phát hiện (greeting, recommend, compare...)'],
        [7, 'SuggestedCardIds', 'array<string>', 'Danh sách Id thẻ được gợi ý trong phản hồi'],
        [8, 'QuickReplies', 'array<string>', 'Danh sách gợi ý nhanh trả về cho user'],
        [9, 'ResponseTimeMs', 'int', 'Thời gian phản hồi (mili-giây)'],
        [10, 'CreatedAt', 'DateTime', 'Thời gian tạo tin nhắn'],
    ]
)

# Save
output_path = '/Users/sivan/chitieuthongminh/database_documentation.docx'
doc.save(output_path)
print(f'✅ Saved to {output_path}')
