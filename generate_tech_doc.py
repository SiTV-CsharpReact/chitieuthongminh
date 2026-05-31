from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)

# ===== TITLE =====
title = doc.add_heading('TỔNG QUAN CÔNG NGHỆ SỬ DỤNG TRONG DỰ ÁN', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Dự án: CredBack – Nền tảng tư vấn thẻ tín dụng thông minh')
run.font.size = Pt(13)
run.italic = True

# ===== 1. GIỚI THIỆU =====
doc.add_heading('1. Giới thiệu chung', level=2)
doc.add_paragraph(
    'CredBack là nền tảng web tư vấn và so sánh thẻ tín dụng thông minh, '
    'giúp người dùng Việt Nam tìm được thẻ tín dụng phù hợp nhất dựa trên '
    'thu nhập, thói quen chi tiêu và nhu cầu cá nhân. Hệ thống được xây dựng '
    'theo kiến trúc Client-Server với Frontend (Next.js), Backend API (ASP.NET Core) '
    'và Database (MongoDB), tích hợp module AI Chatbot và Web Scraper tự động.'
)

# ===== 2. KIẾN TRÚC =====
doc.add_heading('2. Kiến trúc hệ thống', level=2)

arch_table = doc.add_table(rows=5, cols=3)
arch_table.style = 'Light Grid Accent 1'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tầng', 'Công nghệ', 'Vai trò']
for i, h in enumerate(headers):
    cell = arch_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

arch_data = [
    ('Frontend', 'Next.js 16, React 19, TypeScript 5, Tailwind CSS 4, Shadcn UI', 'Giao diện người dùng, SSR/SSG, responsive design'),
    ('Backend', 'ASP.NET Core 8 (C#), REST API, Swagger', 'Xử lý logic nghiệp vụ, xác thực, cung cấp API'),
    ('Database', 'MongoDB, Docker Compose', 'Lưu trữ dữ liệu NoSQL linh hoạt'),
    ('AI & Algorithms', 'ID3 Decision Tree, NLP Chatbot, Scoring Engine', 'Gợi ý thẻ thông minh, chatbot tư vấn tài chính'),
]
for idx, (layer, tech, role) in enumerate(arch_data):
    row = arch_table.rows[idx + 1]
    row.cells[0].text = layer
    row.cells[1].text = tech
    row.cells[2].text = role

# ===== 3. CHI TIẾT =====
doc.add_heading('3. Chi tiết công nghệ', level=2)

techs = [
    {
        'name': 'Next.js 16 + React 19',
        'cat': 'Frontend Framework',
        'desc': 'Framework React thế hệ mới với App Router, Server Components và Turbopack. '
                'Hỗ trợ SSR/SSG giúp tối ưu SEO và tốc độ tải trang.',
        'features': [
            'App Router với Layout lồng nhau (nested layouts)',
            'Server & Client Components',
            'Dynamic routing cho trang chi tiết thẻ tín dụng',
            'Tối ưu SEO tự động với metadata API',
        ],
    },
    {
        'name': 'TypeScript 5',
        'cat': 'Ngôn ngữ lập trình Frontend',
        'desc': 'Superset của JavaScript với hệ thống kiểu tĩnh mạnh mẽ, giúp phát hiện lỗi '
                'sớm trong quá trình phát triển, tăng khả năng bảo trì code.',
        'features': [
            'Type-safe API calls với interface định nghĩa rõ ràng',
            'Generic types cho các component tái sử dụng',
            'Strict null checks giảm lỗi runtime',
            'IDE auto-completion nâng cao năng suất',
        ],
    },
    {
        'name': 'Tailwind CSS 4 + Shadcn UI',
        'cat': 'Styling & UI Components',
        'desc': 'Utility-first CSS framework kết hợp component library Shadcn UI tạo giao diện '
                'premium với dark mode, glassmorphism và micro-animations.',
        'features': [
            'Design System với CSS custom properties (variables)',
            'Dark mode tự động chuyển đổi',
            'Responsive hoàn toàn trên mọi thiết bị',
            'Component library Shadcn UI cho Admin CMS',
        ],
    },
    {
        'name': 'ASP.NET Core 8 (C#)',
        'cat': 'Backend Framework',
        'desc': 'Web API hiệu năng cao với kiến trúc RESTful, hỗ trợ Dependency Injection, '
                'JWT Authentication và Swagger documentation tự động.',
        'features': [
            'RESTful API với 13 Controllers',
            'JWT Bearer Authentication & Google OAuth 2.0',
            'Swagger UI tự động sinh tài liệu API',
            'Dependency Injection cho 8 Services',
        ],
    },
    {
        'name': 'MongoDB + Docker',
        'cat': 'Cơ sở dữ liệu',
        'desc': 'NoSQL Document Database linh hoạt, triển khai qua Docker Compose. '
                'Lưu trữ dữ liệu thẻ tín dụng, người dùng, chi tiêu và lịch sử chat.',
        'features': [
            'Document-oriented storage với schema linh hoạt',
            'Docker Compose deployment đơn giản',
            'MongoDB Driver 3.7 cho .NET',
            'Collections: users, credit_cards, categories, articles, chat_logs, spendings',
        ],
    },
    {
        'name': 'NLP Chatbot Engine',
        'cat': 'AI Chatbot',
        'desc': 'Chatbot tư vấn tài chính thông minh với 11 intent detection patterns, '
                'hỗ trợ tiếng Việt. Tự động gợi ý thẻ, so sánh và tra cứu theo ngữ cảnh.',
        'features': [
            'Regex-based Intent Detection (11 intents: greeting, recommend, compare, card_info, bank_search, cashback, salary, annual_fee, top_cards, count, help)',
            'Multi-turn Conversation với lịch sử chat',
            'Card Recommendation Engine tích hợp',
            'Quick Reply Suggestions cho UX tốt hơn',
        ],
    },
    {
        'name': 'Thuật toán ID3 Decision Tree',
        'cat': 'Thuật toán gợi ý',
        'desc': 'Thuật toán cây quyết định ID3 phân tích đa tiêu chí để xếp hạng và gợi ý '
                'thẻ tín dụng phù hợp nhất cho từng người dùng.',
        'features': [
            'Category Matching với trọng số (ăn uống, mua sắm, du lịch, xăng dầu...)',
            'Salary Bracket Scoring (perfect match, over-qualified, basic)',
            'Income Level Analysis (High/Low vs phí thường niên)',
            'Credit Score Ranking',
        ],
    },
    {
        'name': 'Bank Card Web Scraper',
        'cat': 'Web Scraping & Automation',
        'desc': 'Hệ thống thu thập dữ liệu thẻ tín dụng tự động từ 17+ ngân hàng Việt Nam. '
                'Sử dụng HtmlAgilityPack để parse HTML và trích xuất thông tin chi tiết.',
        'features': [
            'Hỗ trợ 17 ngân hàng: VIB, HSBC, Techcombank, VPBank, ACB, BIDV, MB Bank, Sacombank, TPBank, Shinhan, OCB, Eximbank, HDBank, Vietcombank, UOB, MSB, Standard Chartered',
            'HTML parsing với HtmlAgilityPack',
            'Next.js SSR data extraction (__NEXT_DATA__)',
            'Auto-enrich từ detail pages (parallel fetch)',
        ],
    },
    {
        'name': 'JWT + Google OAuth 2.0',
        'cat': 'Xác thực & Bảo mật',
        'desc': 'Hệ thống xác thực đa lớp: JWT Bearer tokens cho API, Google OAuth 2.0 cho '
                'đăng nhập nhanh, BCrypt cho mã hóa mật khẩu.',
        'features': [
            'JWT Token với thời hạn 7 ngày',
            'Google OAuth 2.0 Integration',
            'BCrypt password hashing (BCrypt.Net-Next)',
            'Role-based Authorization (Admin/User)',
        ],
    },
    {
        'name': 'Recharts',
        'cat': 'Data Visualization',
        'desc': 'Thư viện biểu đồ React cho phân tích chi tiêu trực quan với nhiều loại chart.',
        'features': [
            'Biểu đồ chi tiêu theo tháng/danh mục',
            'Interactive tooltips và responsive layout',
        ],
    },
    {
        'name': 'TinyMCE React',
        'cat': 'Rich Text Editor',
        'desc': 'WYSIWYG editor tích hợp cho module quản lý bài viết trong Admin CMS.',
        'features': [
            'Rich text formatting và image embedding',
            'HTML output cho SEO-friendly articles',
        ],
    },
    {
        'name': 'File Manager System',
        'cat': 'Quản lý Media',
        'desc': 'Hệ thống quản lý ảnh tích hợp với upload, tìm kiếm, di chuyển file.',
        'features': [
            'Upload, rename, move, delete files/folders',
            'Vietnamese filename sanitization (slug)',
            'Storage monitoring (giới hạn 5GB)',
        ],
    },
]

for t in techs:
    doc.add_heading(f'{t["name"]} ({t["cat"]})', level=3)
    doc.add_paragraph(t['desc'])
    for f in t['features']:
        doc.add_paragraph(f, style='List Bullet')

# ===== 4. BẢNG TỔNG HỢP =====
doc.add_heading('4. Bảng tổng hợp thư viện & package', level=2)

pkg_table = doc.add_table(rows=1, cols=4)
pkg_table.style = 'Light Grid Accent 1'
pkg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Thành phần', 'Tên Package', 'Phiên bản', 'Mục đích']):
    pkg_table.rows[0].cells[i].text = h
    pkg_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

packages = [
    ('Frontend', 'next', '16.2.3', 'React framework SSR/SSG'),
    ('Frontend', 'react / react-dom', '19.2.4', 'UI library'),
    ('Frontend', 'typescript', '5.x', 'Type-safe JavaScript'),
    ('Frontend', 'tailwindcss', '4.x', 'Utility-first CSS'),
    ('Frontend', 'shadcn', '4.4.0', 'UI component library'),
    ('Frontend', 'recharts', '3.8.1', 'Biểu đồ phân tích'),
    ('Frontend', '@tinymce/tinymce-react', '6.3.0', 'Rich text editor'),
    ('Frontend', '@react-oauth/google', '0.13.5', 'Google OAuth login'),
    ('Frontend', 'axios', '1.15.0', 'HTTP client'),
    ('Frontend', 'lucide-react', '1.8.0', 'Icon library'),
    ('Frontend', 'qr-code-styling', '1.9.2', 'QR code generator'),
    ('Backend', 'ASP.NET Core', '8.0', 'Web API framework'),
    ('Backend', 'MongoDB.Driver', '3.7.1', 'MongoDB connector'),
    ('Backend', 'BCrypt.Net-Next', '4.1.0', 'Password hashing'),
    ('Backend', 'Google.Apis.Auth', '1.73.0', 'Google OAuth verify'),
    ('Backend', 'HtmlAgilityPack', '1.12.4', 'HTML parser (scraper)'),
    ('Backend', 'JwtBearer Auth', '8.0.0', 'JWT authentication'),
    ('Backend', 'Swashbuckle', '6.4.0', 'Swagger/OpenAPI docs'),
    ('Database', 'MongoDB', 'latest', 'NoSQL database'),
    ('Infra', 'Docker Compose', '-', 'Container orchestration'),
]
for comp, name, ver, purpose in packages:
    row = pkg_table.add_row()
    row.cells[0].text = comp
    row.cells[1].text = name
    row.cells[2].text = ver
    row.cells[3].text = purpose

# ===== 5. SỐ LIỆU =====
doc.add_heading('5. Quy mô dự án', level=2)
stats = [
    ('Số ngân hàng hỗ trợ scraping', '17'),
    ('Số API Controllers', '13'),
    ('Số Backend Services', '8'),
    ('Số AI Chatbot Intents', '11'),
    ('Số modules Frontend (routes)', '8 client + 7 admin'),
    ('Frontend packages', '12 dependencies'),
    ('Backend NuGet packages', '7 packages'),
]
st = doc.add_table(rows=len(stats)+1, cols=2)
st.style = 'Light Grid Accent 1'
st.alignment = WD_TABLE_ALIGNMENT.CENTER
st.rows[0].cells[0].text = 'Chỉ số'
st.rows[0].cells[1].text = 'Giá trị'
for c in st.rows[0].cells:
    c.paragraphs[0].runs[0].bold = True
for i, (k, v) in enumerate(stats):
    st.rows[i+1].cells[0].text = k
    st.rows[i+1].cells[1].text = v

# Save
out = '/Users/sivan/chitieuthongminh/tong_quan_cong_nghe.docx'
doc.save(out)
print(f'Saved: {out}')
